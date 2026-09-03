from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path('/Users/jc/Desktop/work/peixun/output/首周任务：完成从“新同学”到“团队成员”的第一步.docx')
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_run_font(run, size=11, color=INK, bold=False):
    run.font.name = 'Arial Unicode MS'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_section(doc, number, title, intro, items):
    p = doc.add_paragraph(style='Heading 2')
    add_text(p, f'{number}. {title}', size=13, color=BLUE, bold=True)
    p.paragraph_format.keep_with_next = True
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_text(p, intro, color=MUTED)
    for item in items:
        if isinstance(item, tuple):
            lead, body = item
            p = add_bullet(doc, '')
            add_text(p, lead, bold=True)
            add_text(p, body)
        else:
            add_bullet(doc, item)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial Unicode MS'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 12, 7),
        ('Heading 2', 13, BLUE, 10, 5),
        ('Heading 3', 11.5, INK, 7, 3),
    ]:
        style = styles[name]
        style.font.name = 'Arial Unicode MS'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    add_text(title, '首周任务：完成从“新同学”到“团队成员”的第一步', size=20, color=BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, '新人入职行动指南｜第一周使用', size=11, color=MUTED)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    add_text(intro, '本手册帮助新同事在入职第一周完成基础配置、了解公司与产品、加入协作网络，并建立对项目和工作方式的初步认识。', size=11, color=INK)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ['本周目标', '完成标准']):
        set_cell_shading(cell, 'EAF4FB')
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, label, size=10.5, color=BLUE, bold=True)
    row = table.add_row().cells
    add_text(row[0].paragraphs[0], '知道找谁、去哪里、用什么工具完成工作。')
    add_text(row[1].paragraphs[0], '完成下方六项行动，并与导师或直属负责人确认下一步。')

    h = doc.add_paragraph(style='Heading 1')
    add_text(h, '入职前半程：先把环境和协作网络建立起来', size=16, color=BLUE, bold=True)
    add_section(doc, 1, '完成入职手续及办公环境配置', '先确保账号、设备和权限可用，再开始正式学习。', [
        ('人事手续：', '与 HR（Henna）确认合同、岗位说明及入职资料。'),
        ('账号与权限：', '登录飞书，按需申请邮箱、项目文档及其他工作系统权限。'),
        ('设备与软件：', '与 IT（小贾、鲁班）确认电脑、网络、耳机及岗位所需软件正常。'),
        '发现权限或设备问题时，记录具体现象、截图和影响范围，及时反馈给对应负责人。',
    ])
    add_section(doc, 2, '了解公司与体验产品', '建立对公司、业务和日常工作规则的基本认识。', [
        ('公司与业务：', '了解公司发展、组织结构、产品矩阵及当前重点方向。'),
        ('新人培训：', '完成新人培训课程和测评，遇到不理解的内容及时记录问题。'),
        ('办公规则：', '了解考勤、请假、报销、会议室使用和办公安全要求。'),
        ('产品体验：', '下载并体验 Bingo Voyage、Bingo Frenzy、Matching Story 等产品，记录初步感受。'),
        ('工作同步：', '按团队要求提交日报，逐步熟悉信息同步和反馈节奏。'),
    ])
    add_section(doc, 3, '加入团队及项目协作群', '进入正确的信息流，避免遗漏通知和项目上下文。', [
        ('加入群组：', '加入部门群、项目群、新人群及岗位相关群组，完善个人资料。'),
        ('阅读背景：', '查看群公告、置顶消息、项目介绍和近期通知，了解群组用途。'),
        ('建立通讯录：', '记下直属负责人、导师、项目负责人及关键协作者的联系方式。'),
    ])
    add_section(doc, 4, '认识团队成员及关键协作伙伴', '通过真实交流理解角色分工和协作方式。', [
        ('直属负责人：', '与直属负责人进行 1:1，确认岗位期待、近期重点和沟通方式。'),
        ('导师与团队：', '了解导师、同组成员的职责边界，以及日常协作节奏。'),
        ('跨团队伙伴：', '认识产品、研发、设计、运营或其他会共同推进工作的伙伴。'),
    ])

    h = doc.add_paragraph(style='Heading 1')
    add_text(h, '入职行动：第一周完成这六件事', size=16, color=BLUE, bold=True)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(7)
    add_text(intro, '以下行动是第一周的可执行清单。完成一项就记录结果、问题和后续动作。', color=MUTED)
    actions = [
        ('完成入职手续及办公环境配置', '确认飞书、邮箱、办公设备及岗位所需权限可以正常使用。'),
        ('了解公司及体验产品', '完成新人培训和测评，并对公司产品形成初步体验记录。'),
        ('加入团队及项目协作群', '加入必要群组，阅读关键资料，知道信息应该在哪里同步。'),
        ('认识团队成员及关键协作伙伴', '完成与直属负责人、导师及关键协作者的初次沟通。'),
        ('了解项目背景与当前进展', '阅读项目介绍、版本计划和新人材料，了解目标、阶段、里程碑与主要挑战。'),
        ('完成首周复盘', '总结完成事项、已掌握内容、疑问与支持需求，并确认下一周计划。'),
    ]
    for title_text, body in actions:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        add_text(p, title_text, size=11, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, body, color=MUTED)

    h = doc.add_paragraph(style='Heading 1')
    add_text(h, '完成后的复盘提示', size=16, color=BLUE, bold=True)
    for item in [
        '哪些事项已经完成，哪些事项仍需要他人协助？',
        '我是否知道常用工具、流程和信息入口？',
        '我对团队目标、项目背景和自己的下一步工作还有哪些疑问？',
        '与导师或直属负责人确认：下周优先级、交付方式和跟进时间。',
    ]:
        add_bullet(doc, item)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    add_text(note, '说明：第一周不要求立即产出复杂成果。遇到问题先提问、记录和反馈，团队会提供必要的资料与支持。', size=10, color=MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, '乐堡家园｜新人培训手册', size=9, color=MUTED)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()
